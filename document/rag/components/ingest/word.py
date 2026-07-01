import io
import logging
import re
from typing import Optional, Dict, Any, BinaryIO, List
from xml.etree import ElementTree

from core.ports.ingest import (
    IngestPort,
    IngestResult,
    IngestConfig,
    IngestStatus,
    DocumentFormat,
)


class WordIngestAdapter:
    """
    æ¯ææ ¼å¼ï¼?    - .docx (Office Open XML)
    - .doc (æ§æ ¼å¼ï¼éè¦?python-docx2txt æ?antiword)
    
    ç¹æ§ï¼
    - æ é¢ãæ®µè½ãåè¡¨æå?    - è¡¨æ ¼æå
    - å¾çæåï¼base64 æé¾æ¥ï¼
    - æ ·å¼ä¿¡æ¯ä¿ç
    """
    
    def __init__(
        self,
        extract_styles: bool = False,
        merge_list_items: bool = True
    ):
        self._logger = logging.getLogger("ingest.word")
        self._extract_styles = extract_styles
        self._merge_list_items = merge_list_items
    
    def ingest(
        self,
        source: BinaryIO,
        doc_format: DocumentFormat,
        doc_id: str,
        config: Optional[IngestConfig] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> IngestResult:
        config = config or IngestConfig()
        metadata = metadata or {}
        
        metadata["doc_id"] = doc_id
        metadata["doc_format"] = "word"
        
        try:
            return self._extract_docx(source, config, metadata)
        except (RuntimeError, OSError, ValueError) as e:
            self._logger.error(f"Word extraction failed: {e}")
            
            if config.ocr_fallback:
                self._logger.info("Attempting OCR fallback for .doc")
                return self._fallback_doc_extraction(source, config, metadata)
            
            return IngestResult(
                content="",
                metadata=metadata,
                status=IngestStatus.FAILED,
                errors=[str(e)]
            )
    
    def ingest_from_path(
        self,
        file_path: str,
        doc_id: str,
        config: Optional[IngestConfig] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> IngestResult:
        metadata = metadata or {}
        metadata["source_path"] = file_path
        
        with open(file_path, "rb") as f:
            return self.ingest(f, DocumentFormat.WORD, doc_id, config, metadata)
    
    def supports_format(self, doc_format: DocumentFormat) -> bool:
        return doc_format == DocumentFormat.WORD
    
    def _extract_docx(
        self,
        source: BinaryIO,
        config: IngestConfig,
        metadata: Dict[str, Any]
    ) -> IngestResult:
        from docx import Document
        
        source.seek(0)
        doc = Document(source)
        
        sections = []
        all_tables = []
        all_images = []
        current_section = []
        current_heading = ""
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            style_name = para.style.name if para.style else ""
            
            if "Heading" in style_name or style_name.startswith("æ é¢"):
                if current_section:
                    sections.append({
                        "heading": current_heading,
                        "content": "\n".join(current_section),
                    })
                    current_section = []
                current_heading = text
            else:
                if self._merge_list_items and style_name.startswith("List"):
                    text = f"- {text}"
                current_section.append(text)
        
        if current_section:
            sections.append({
                "heading": current_heading,
                "content": "\n".join(current_section),
            })
        
        if config.extract_tables:
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    all_tables.append({
                        "index": table_idx,
                        "data": table_data,
                        "rows": len(table_data),
                        "cols": len(table_data[0]) if table_data else 0,
                    })
        
        if config.extract_images:
            try:
                images = self._extract_images_from_docx(source)
                all_images.extend(images)
            except (OSError, RuntimeError, ValueError) as e:
                self._logger.warning(f"Image extraction failed: {e}")
        
        full_content_parts = []
        for section in sections:
            if section["heading"]:
                full_content_parts.append(f"## {section['heading']}")
            full_content_parts.append(section["content"])
        
        full_content = "\n\n".join(full_content_parts)
        
        metadata["section_count"] = len(sections)
        metadata["table_count"] = len(all_tables)
        metadata["image_count"] = len(all_images)
        
        pages = [{
            "page_num": 1,
            "content": full_content,
            "char_count": len(full_content),
        }]
        
        return IngestResult(
            content=full_content,
            metadata=metadata,
            status=IngestStatus.SUCCESS,
            pages=pages,
            tables=all_tables,
            images=all_images,
        )
    
    def _extract_images_from_docx(self, source: BinaryIO) -> List[Dict[str, Any]]:
        import zipfile
        
        images = []
        source.seek(0)
        
        with zipfile.ZipFile(source) as zf:
            for name in zf.namelist():
                if name.startswith('word/media/'):
                    image_data = zf.read(name)
                    ext = name.split('.')[-1].lower()
                    
                    images.append({
                        "filename": name.split('/')[-1],
                        "format": ext,
                        "size": len(image_data),
                    })
        
        return images
    
    def _fallback_doc_extraction(
        self,
        source: BinaryIO,
        config: IngestConfig,
        metadata: Dict[str, Any]
    ) -> IngestResult:
        errors = []
        
        try:
            import docx2txt
            
            source.seek(0)
            content = docx2txt.process(source)
            
            metadata["backend"] = "docx2txt"
            
            return IngestResult(
                content=content,
                metadata=metadata,
                status=IngestStatus.PARTIAL,
                errors=["Used docx2txt fallback"],
            )
        except ImportError:
            errors.append("docx2txt not available")
        
        try:
            import subprocess
            source.seek(0)
            
            temp_path = f"/tmp/{metadata.get('doc_id', 'temp')}.doc"
            with open(temp_path, 'wb') as f:
                f.write(source.read())
            
            result = subprocess.run(
                ['antiword', temp_path],
                capture_output=True,
                text=True
            )
            
            if result.returncode == 0:
                metadata["backend"] = "antiword"
                return IngestResult(
                    content=result.stdout,
                    metadata=metadata,
                    status=IngestStatus.PARTIAL,
                    errors=["Used antiword fallback"],
                )
        except (OSError, RuntimeError, ValueError) as e:
            errors.append(f"antiword failed: {e}")
        
        return IngestResult(
            content="",
            metadata=metadata,
            status=IngestStatus.FAILED,
            errors=errors,
        )


class WordToJsonAdapter:
    """
    è¾åºç»æï¼?    {
        "title": "ææ¡£æ é¢",
        "sections": [
            {
                "heading": "ç« èæ é¢",
                "content": "ç« èåå®¹",
                "level": 1
            }
        ],
        "tables": [...],
        "metadata": {...}
    }
    """
    
    def __init__(self):
        self._logger = logging.getLogger("ingest.word_json")
    
    def convert(
        self,
        source: BinaryIO,
        config: Optional[IngestConfig] = None
    ) -> Dict[str, Any]:
        from docx import Document
        
        source.seek(0)
        doc = Document(source)
        
        result = {
            "title": "",
            "sections": [],
            "tables": [],
            "metadata": {},
        }
        
        title_found = False
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            style_name = para.style.name if para.style else ""
            
            if not title_found and (style_name == "Title" or "Heading 1" in style_name):
                result["title"] = text
                title_found = True
                continue
            
            heading_match = re.match(r"Heading (\d+)", style_name)
            if heading_match or style_name.startswith("æ é¢"):
                level = 1
                if heading_match:
                    level = int(heading_match.group(1))
                
                result["sections"].append({
                    "heading": text,
                    "content": "",
                    "level": level,
                })
            elif result["sections"]:
                result["sections"][-1]["content"] += text + "\n"
            elif not title_found:
                result["sections"].append({
                    "heading": "",
                    "content": text,
                    "level": 0,
                })
        
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            result["tables"].append(table_data)
        
        result["metadata"] = {
            "section_count": len(result["sections"]),
            "table_count": len(result["tables"]),
        }
        
        return result
    
    def convert_from_path(self, file_path: str) -> Dict[str, Any]:
        with open(file_path, "rb") as f:
            return self.convert(f)
