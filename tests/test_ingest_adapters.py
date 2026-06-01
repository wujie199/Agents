import io
import pytest
from PIL import Image

from core.ports.ingest import (
    IngestConfig,
    IngestStatus,
    DocumentFormat,
)
from document.rag.pipeline.ingest.adapters.word_adapter import WordIngestAdapter
from document.rag.pipeline.ingest.adapters.layout_ocr_adapter import LayoutOCRAdapter
from document.rag.pipeline.ingest.adapters.simplified_adapter import (
    SimplifiedIngestAdapter,
    SimplifiedIngestPipeline,
    build_simplified_ingest_adapter,
)


class TestWordIngestAdapter:
    
    @pytest.fixture
    def word_adapter(self):
        return WordIngestAdapter()
    
    def test_supports_format(self, word_adapter):
        assert word_adapter.supports_format(DocumentFormat.WORD) is True
        assert word_adapter.supports_format(DocumentFormat.PDF) is False
    
    def test_ingest_simple_docx(self, word_adapter):
        try:
            from docx import Document
            
            doc = Document()
            doc.add_heading("Test Document", 0)
            doc.add_paragraph("First paragraph content.")
            doc.add_paragraph("Second paragraph content.")
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            result = word_adapter.ingest(
                buffer,
                DocumentFormat.WORD,
                "test_doc",
                IngestConfig()
            )
            
            assert result.status == IngestStatus.SUCCESS
            assert "First paragraph" in result.content
            assert result.metadata["doc_id"] == "test_doc"
            
        except ImportError:
            pytest.skip("python-docx not available")
    
    def test_ingest_with_tables(self, word_adapter):
        try:
            from docx import Document
            
            doc = Document()
            doc.add_paragraph("Document with table")
            
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Header 1"
            table.cell(0, 1).text = "Header 2"
            table.cell(1, 0).text = "Data 1"
            table.cell(1, 1).text = "Data 2"
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            config = IngestConfig(extract_tables=True)
            result = word_adapter.ingest(
                buffer,
                DocumentFormat.WORD,
                "test_doc",
                config
            )
            
            assert result.status == IngestStatus.SUCCESS
            assert len(result.tables) == 1
            assert result.tables[0]["rows"] == 2
            
        except ImportError:
            pytest.skip("python-docx not available")


class TestLayoutOCRAdapter:
    
    @pytest.fixture
    def layout_adapter(self):
        return LayoutOCRAdapter(ocr_backend="auto")
    
    def test_supports_format(self, layout_adapter):
        assert layout_adapter.supports_format(DocumentFormat.PDF) is True
        assert layout_adapter.supports_format(DocumentFormat.IMAGE) is True
        assert layout_adapter.supports_format(DocumentFormat.HTML) is True
        assert layout_adapter.supports_format(DocumentFormat.WORD) is False
    
    def test_detect_ocr_backend(self, layout_adapter):
        assert layout_adapter._ocr_backend in ["paddleocr", "tesseract", "easyocr", "none"]
    
    def test_ingest_image(self, layout_adapter):
        if layout_adapter._ocr_backend == "none":
            pytest.skip("No OCR backend available")
        
        img = Image.new('RGB', (200, 50), color='white')
        
        buffer = io.BytesIO()
        img.save(buffer, format='PNG')
        buffer.seek(0)
        
        result = layout_adapter.ingest(
            buffer,
            DocumentFormat.IMAGE,
            "test_image",
            IngestConfig()
        )
        
        assert result.status in [IngestStatus.SUCCESS, IngestStatus.FAILED]
        assert result.metadata["doc_format"] == "image"
        assert result.metadata["ocr_backend"] == layout_adapter._ocr_backend
    
    def test_pdf_to_images(self, layout_adapter):
        try:
            import fitz
            
            doc = fitz.open()
            page = doc.new_page(width=612, height=792)
            page.insert_text((100, 100), "Test PDF page content")
            
            buffer = io.BytesIO()
            doc.save(buffer)
            doc.close()
            buffer.seek(0)
            
            config = IngestConfig(dpi=150)
            images = layout_adapter._pdf_to_images(buffer, config)
            
            assert len(images) == 1
            assert images[0][1] == 1
            
        except ImportError:
            pytest.skip("PyMuPDF not available")


class TestSimplifiedIngestAdapter:
    
    @pytest.fixture
    def simplified_adapter(self):
        return build_simplified_ingest_adapter()
    
    def test_detect_format(self, simplified_adapter):
        assert simplified_adapter._detect_format("test.docx") == DocumentFormat.WORD
        assert simplified_adapter._detect_format("test.doc") == DocumentFormat.WORD
        assert simplified_adapter._detect_format("test.pdf") == DocumentFormat.PDF
        assert simplified_adapter._detect_format("test.png") == DocumentFormat.IMAGE
        assert simplified_adapter._detect_format("test.html") == DocumentFormat.HTML
        assert simplified_adapter._detect_format("test.unknown") is None
    
    def test_get_processing_mode(self, simplified_adapter):
        assert simplified_adapter.get_processing_mode(DocumentFormat.WORD) == "word_structured"
        assert simplified_adapter.get_processing_mode(DocumentFormat.PDF) == "layout_ocr"
        assert simplified_adapter.get_processing_mode(DocumentFormat.IMAGE) == "layout_ocr"
        assert simplified_adapter.get_processing_mode(DocumentFormat.HTML) == "layout_ocr"
    
    def test_ingest_word_document(self, simplified_adapter):
        try:
            from docx import Document
            
            doc = Document()
            doc.add_paragraph("Word document for simplified adapter test")
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            result = simplified_adapter.ingest(
                buffer,
                DocumentFormat.WORD,
                "test_word",
                IngestConfig()
            )
            
            assert result.status == IngestStatus.SUCCESS
            assert "Word document" in result.content
            assert result.metadata["doc_format"] == "word"
            
        except ImportError:
            pytest.skip("python-docx not available")
    
    def test_ingest_image_with_ocr(self, simplified_adapter):
        if simplified_adapter._layout_ocr_adapter._ocr_backend == "none":
            pytest.skip("No OCR backend available")
        
        img = Image.new('RGB', (100, 30), color='white')
        
        buffer = io.BytesIO()
        img.save(buffer, format='PNG')
        buffer.seek(0)
        
        result = simplified_adapter.ingest(
            buffer,
            DocumentFormat.IMAGE,
            "test_image",
            IngestConfig()
        )
        
        assert result.status in [IngestStatus.SUCCESS, IngestStatus.FAILED]
        assert result.metadata["ocr_backend"] is not None


class TestSimplifiedIngestPipeline:
    
    @pytest.fixture
    def pipeline(self):
        return SimplifiedIngestPipeline(
            ingest_adapter=build_simplified_ingest_adapter()
        )
    
    def test_pipeline_word_document(self, pipeline):
        try:
            from docx import Document
            
            doc = Document()
            doc.add_paragraph("Pipeline test with Word document")
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            result = pipeline.ingest(
                buffer,
                DocumentFormat.WORD,
                "test_doc",
                IngestConfig()
            )
            
            assert result.status == IngestStatus.SUCCESS
            assert "char_count" in result.metadata
            assert "word_count" in result.metadata
            
        except ImportError:
            pytest.skip("python-docx not available")
    
    def test_pipeline_with_cleaner(self):
        from document.rag.bridges.composite_cleaner import CleanerAdapter
        
        adapter = build_simplified_ingest_adapter()
        cleaner = CleanerAdapter()
        
        pipeline = SimplifiedIngestPipeline(
            ingest_adapter=adapter,
            cleaner_adapter=cleaner
        )
        
        try:
            from docx import Document
            
            doc = Document()
            doc.add_paragraph("Test  content  with  extra  spaces")
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            result = pipeline.ingest(
                buffer,
                DocumentFormat.WORD,
                "test_doc",
                IngestConfig()
            )
            
            assert result.status == IngestStatus.SUCCESS
            
        except ImportError:
            pytest.skip("python-docx not available")


class TestIngestResult:
    
    def test_ingest_result_properties(self):
        from core.ports.ingest import IngestResult
        
        result = IngestResult(
            content="Test content",
            pages=[{"page_num": 1}, {"page_num": 2}],
            tables=[{"data": []}],
            images=[{"format": "png"}],
        )
        
        assert result.page_count == 2
        assert result.has_tables is True
        assert result.has_images is True
    
    def test_ingest_result_defaults(self):
        from core.ports.ingest import IngestResult
        
        result = IngestResult(content="Test")
        
        assert result.page_count == 0
        assert result.has_tables is False
        assert result.has_images is False
        assert result.status == IngestStatus.SUCCESS
        assert result.errors == []
